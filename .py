from docx import Document

# Create a new Word document
doc = Document()

# Title
doc.add_heading("Proposal for Dedicated Mentorship Program 2025", level=1)

# Project Details
doc.add_heading("Project Title:", level=2)
doc.add_paragraph("National Livelihood Dashboard – eGovernments Foundation")

doc.add_heading("Track:", level=2)
doc.add_paragraph("Coding & AI")

# Problem Statement
doc.add_heading("1. Problem Statement", level=2)
doc.add_paragraph(
    "India's public employment schemes and livelihood programs generate vast amounts of data, "
    "but these datasets are often siloed and not available in a consolidated, actionable format. "
    "Policymakers and government officers need access to real-time insights across states and districts "
    "to evaluate outcomes and adjust implementation strategies. A centralized dashboard can bridge this gap "
    "by visualizing key livelihood indicators such as job creation, skills training, and program efficiency."
)

# Proposed Solution
doc.add_heading("2. Proposed Solution", level=2)
doc.add_paragraph(
    "This project proposes building a cloud-hosted, interactive dashboard that integrates various employment-related datasets "
    "and presents them in an accessible and intuitive manner. The dashboard will:\n"
    "- Provide KPIs such as total jobs created, active employment schemes, regional breakdowns.\n"
    "- Enable filters by state, district, and time periods.\n"
    "- Include visualizations (charts, graphs, maps) to support data analysis.\n"
    "- Offer secure, role-based access for policymakers and administrators.\n"
    "- Be scalable for additional datasets or features in the future."
)

# Tech Stack
doc.add_heading("3. Tech Stack", level=2)
doc.add_paragraph(
    "- Backend: Python (FastAPI or Flask), PostgreSQL/SQLite\n"
    "- Frontend: React.js, Chart.js or Plotly.js for data visualization\n"
    "- Data Handling: REST APIs, scheduled ETL scripts (Python), data.gov.in APIs\n"
    "- Deployment: Docker, GitHub Actions, Railway/Render"
)

# Timeline
doc.add_heading("4. Timeline (12 Weeks)", level=2)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Week'
hdr_cells[1].text = 'Goals'
weeks = [
    ("1–2", "Understand the problem scope, collect mock datasets, define KPIs"),
    ("3–4", "Develop backend APIs, design database schema"),
    ("5–6", "Start frontend dashboard with sample data visualizations"),
    ("7–8", "Add interactivity: filters, region selection, data export options"),
    ("9–10", "Deploy using Docker and CI/CD tools"),
    ("11–12", "Testing, documentation, and final feedback incorporation")
]
for week, goal in weeks:
    row_cells = table.add_row().cells
    row_cells[0].text = week
    row_cells[1].text = goal

# Past Experience
doc.add_heading("5. Past Experience", level=2)
doc.add_paragraph(
    "I have developed a Cloud-Based Inventory Management System using Python and SQL. Key features include:\n"
    "- Inventory tracking and reporting\n"
    "- Real-time CRUD operations via REST API\n"
    "- SQLite database integration\n"
    "- Modular design with separate billing and admin components"
)

# Contribution Plan
doc.add_heading("6. Contribution Plan", level=2)
doc.add_paragraph(
    "- Build and test backend APIs for data ingestion and retrieval.\n"
    "- Create responsive frontend components to display KPIs and charts.\n"
    "- Use publicly available datasets to populate and test the dashboard.\n"
    "- Write clear documentation and user guides.\n"
    "- Actively participate in GitHub discussions and maintain code quality."
)

# Optional Contributions
doc.add_heading("7. GitHub/Design Contributions (Optional)", level=2)
doc.add_paragraph(
    "Once selected, I plan to contribute to:\n"
    "- Creating issue reports and pull requests for dashboard components\n"
    "- Enhancing documentation or frontend styling of related repositories"
)

# Personal Information
doc.add_paragraph("\nName: Devang Sadanand Angchekar")
doc.add_paragraph("Email: devangangchekar2004@gmail.com")
doc.add_paragraph("GitHub: https://github.com/Devang-09")
doc.add_paragraph("Discord: https://githubapp.com/users/1353601265036300440")
doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/devang-angchekar-3583a02b9")
doc.add_paragraph("College: St. Arnold's Night Degree College, University of Mumbai")
doc.add_paragraph("Course: Bachelor of Science in Information Technology (BScIT)")
doc.add_paragraph("Availability: 4 hours/day")

# Save the document
doc.save("C4GT_Proposal_Devang.docx")
